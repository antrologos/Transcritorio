from __future__ import annotations

from pathlib import Path
from typing import Any
import csv
import re

from .config import Paths
from .manifest import selected_rows
from .utils import format_timestamp, read_json, write_json

READABLE_SOFT_PARAGRAPH_CHARS = 850
READABLE_HARD_PARAGRAPH_CHARS = 1200
READABLE_MIN_PARAGRAPH_CHARS = 220
SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[.!?])\s+")
DISPLAY_LABELS = {
    "ENTREVISTADOR": "Entrevistador",
    "ENTREVISTADO": "Entrevistado",
}
FLAG_DISPLAY_LABELS = {
    "inaudivel": "inaud\u00edvel",
    "duvida": "d\u00favida",
    "sobreposicao": "sobreposi\u00e7\u00e3o",
}


def render_outputs(rows: list[dict[str, str]], config: dict, paths: Paths, ids: list[str] | None = None) -> int:
    failures = 0
    speaker_map = read_speaker_map(paths.manifest_dir / "speakers_map.csv")
    for row in selected_rows(rows, ids):
        interview_id = row["interview_id"]
        source_json = find_whisperx_json(paths, interview_id)
        if not source_json:
            print(f"Missing WhisperX JSON for {interview_id}")
            failures += 1
            continue
        data = read_json(source_json)
        diarization_source = "whisperx"
        external_segments = load_external_diarization(paths, interview_id, config)
        if external_segments:
            data = apply_external_diarization(data, external_segments)
            diarization_source = str(config.get("diarization_source"))
        effective_speaker_map = speaker_map_from_labels(data, config.get("speaker_labels"))
        effective_speaker_map.update(speaker_map.get(interview_id, {}))
        canonical = {
            "interview_id": interview_id,
            "source_path": row["source_path"],
            "source_sha256": row.get("source_sha256", ""),
            "asr_model": config["asr_model"],
            "diarization_model": str(config.get("diarize_model", "")) if config.get("diarize") else "",
            "diarization_source": diarization_source,
            "speaker_labels": list(config.get("speaker_labels") or []),
            "turns": build_turns(data, config, effective_speaker_map),
        }
        write_json(paths.canonical_dir / "json" / f"{interview_id}.canonical.json", canonical)
        write_markdown(paths.review_dir / "md" / f"{interview_id}.md", canonical)
        write_nvivo_tsv(paths.asr_dir / "tsv" / f"{interview_id}_nvivo.tsv", canonical)
        write_docx_if_available(paths.review_dir / "docx" / f"{interview_id}.docx", canonical)
    return failures


def build_turns(data: dict[str, Any], config: dict, speaker_map: dict[str, str]) -> list[dict[str, Any]]:
    turns: list[dict[str, Any]] = []
    gap = float(config["turn_gap_seconds"])
    max_turn = float(config["max_turn_seconds"])

    for segment in data.get("segments", []):
        text = " ".join(str(segment.get("text", "")).split())
        if not text:
            continue
        speaker = str(segment.get("speaker", "SPEAKER_UNKNOWN"))
        start = float(segment.get("start", 0) or 0)
        end = float(segment.get("end", start) or start)
        human_label = speaker_map.get(speaker, speaker)

        if turns:
            previous = turns[-1]
            if previous["speaker"] == speaker and start - float(previous["end"]) <= gap and end - float(previous["start"]) <= max_turn:
                previous["end"] = end
                previous["text"] = f"{previous['text']} {text}".strip()
                continue

        turn = {"start": start, "end": end, "speaker": speaker, "human_label": human_label, "text": text}
        if segment.get("needs_speaker_review"):
            turn["flags"] = ["duvida"]
            turn["notes"] = str(segment.get("diarization_review_note") or "Revisar falante.")
        turns.append(turn)
    return turns


def speaker_map_from_labels(data: dict[str, Any], labels: object) -> dict[str, str]:
    if not isinstance(labels, list):
        return {}
    cleaned_labels = [str(label).strip() for label in labels if str(label).strip()]
    if not cleaned_labels:
        return {}
    speakers = sorted_speaker_ids(data)
    return {speaker: cleaned_labels[index] for index, speaker in enumerate(speakers[: len(cleaned_labels)])}


def sorted_speaker_ids(data: dict[str, Any]) -> list[str]:
    seen: list[str] = []
    for segment in data.get("segments", []):
        speaker = str(segment.get("speaker", "")).strip()
        if speaker and speaker not in seen:
            seen.append(speaker)
    return sorted(seen, key=speaker_sort_key)


def speaker_sort_key(label: str) -> tuple[int, int | str]:
    match = re.fullmatch(r"SPEAKER_(\d+)", label)
    if match:
        return (0, int(match.group(1)))
    return (1, label)


def load_external_diarization(paths: Paths, interview_id: str, config: dict) -> list[dict[str, Any]]:
    source = str(config.get("diarization_source", "whisperx")).lower()
    if source not in {"pyannote_exclusive", "exclusive"}:
        return []
    path = paths.diarization_dir / "json" / f"{interview_id}.exclusive.json"
    if not path.exists():
        return []
    payload = read_json(path)
    return [segment for segment in payload.get("segments", []) if isinstance(segment, dict)]


def apply_external_diarization(data: dict[str, Any], diarization_segments: list[dict[str, Any]]) -> dict[str, Any]:
    result = dict(data)
    result["segments"] = []
    for segment in data.get("segments", []):
        if not isinstance(segment, dict):
            continue
        split_segments = split_segment_by_word_diarization(segment, diarization_segments)
        if split_segments:
            result["segments"].extend(split_segments)
            continue
        updated = dict(segment)
        start = float(updated.get("start", 0) or 0)
        end = float(updated.get("end", updated.get("start", 0)) or 0)
        speaker = best_overlap_speaker(start, end, diarization_segments)
        if speaker:
            updated["speaker"] = speaker
        if len(overlapping_speakers(start, end, diarization_segments)) > 1:
            updated["needs_speaker_review"] = True
            updated["diarization_review_note"] = "segmento ASR contem mais de um falante e nao tinha palavras suficientes para dividir"
        result["segments"].append(updated)
    return result


def split_segment_by_word_diarization(segment: dict[str, Any], diarization_segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    words = usable_words(segment.get("words"))
    if len(words) < 2:
        return []
    groups: list[dict[str, Any]] = []
    for word in words:
        speaker = best_overlap_speaker(float(word["start"]), float(word["end"]), diarization_segments) or str(segment.get("speaker", "SPEAKER_UNKNOWN"))
        if not groups or groups[-1]["speaker"] != speaker:
            groups.append({"speaker": speaker, "words": [word]})
        else:
            groups[-1]["words"].append(word)
    if len(groups) < 2:
        return []
    result: list[dict[str, Any]] = []
    for group in groups:
        group_words = group["words"]
        text = join_word_texts(group_words)
        if not text:
            continue
        updated = dict(segment)
        updated["speaker"] = group["speaker"]
        updated["start"] = round(float(group_words[0]["start"]), 3)
        updated["end"] = round(float(group_words[-1]["end"]), 3)
        updated["text"] = text
        updated["words"] = [dict(word["raw"]) for word in group_words]
        updated["diarization_split"] = True
        result.append(updated)
    return result


def usable_words(raw_words: Any) -> list[dict[str, Any]]:
    if not isinstance(raw_words, list):
        return []
    words: list[dict[str, Any]] = []
    for raw in raw_words:
        if not isinstance(raw, dict):
            continue
        text = str(raw.get("word") or raw.get("text") or "").strip()
        if not text:
            continue
        try:
            start = float(raw.get("start"))
            end = float(raw.get("end"))
        except (TypeError, ValueError):
            continue
        if end <= start:
            continue
        words.append({"start": start, "end": end, "text": text, "raw": raw})
    return words


def join_word_texts(words: list[dict[str, Any]]) -> str:
    text = ""
    for word in words:
        piece = str(word.get("text", "")).strip()
        if not piece:
            continue
        if not text:
            text = piece
        elif re.match(r"^[,.;:!?%)]", piece):
            text += piece
        else:
            text += f" {piece}"
    return " ".join(text.split())


def overlapping_speakers(start: float, end: float, diarization_segments: list[dict[str, Any]]) -> set[str]:
    speakers: set[str] = set()
    for segment in diarization_segments:
        diar_start = float(segment.get("start", 0) or 0)
        diar_end = float(segment.get("end", diar_start) or diar_start)
        overlap = max(0.0, min(end, diar_end) - max(start, diar_start))
        if overlap > 0:
            speaker = str(segment.get("speaker", "")).strip()
            if speaker:
                speakers.add(speaker)
    return speakers


def best_overlap_speaker(start: float, end: float, diarization_segments: list[dict[str, Any]]) -> str | None:
    best_speaker = None
    best_overlap = 0.0
    for segment in diarization_segments:
        diar_start = float(segment.get("start", 0) or 0)
        diar_end = float(segment.get("end", diar_start) or diar_start)
        overlap = max(0.0, min(end, diar_end) - max(start, diar_start))
        if overlap > best_overlap:
            best_overlap = overlap
            best_speaker = str(segment.get("speaker", ""))
    return best_speaker if best_speaker and best_overlap > 0 else None


def build_readable_blocks(turns: list[dict[str, Any]]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []

    for turn in turns:
        text = " ".join(str(turn.get("text", "")).split())
        flags = " ".join(f"[{FLAG_DISPLAY_LABELS.get(flag, flag)}]" for flag in turn_flags(turn))
        text = append_text(text, flags)
        if not text:
            continue

        speaker = editorial_speaker_key(turn)
        label = format_speaker_label(str(turn.get("human_label") or speaker))
        if blocks and blocks[-1]["speaker"] == speaker:
            blocks[-1]["end"] = float(turn.get("end", blocks[-1]["end"]) or blocks[-1]["end"])
            blocks[-1]["chunks"].append(text)
            continue

        blocks.append(
            {
                "start": float(turn.get("start", 0) or 0),
                "end": float(turn.get("end", turn.get("start", 0)) or 0),
                "speaker": speaker,
                "label": label,
                "chunks": [text],
            }
        )

    for block in blocks:
        block["paragraphs"] = split_readable_paragraphs(block.pop("chunks"))

    return blocks


def format_speaker_label(label: str) -> str:
    normalized = " ".join(str(label).split())
    return DISPLAY_LABELS.get(normalized.upper(), normalized)


def editorial_speaker_key(turn: dict[str, Any]) -> str:
    label = str(turn.get("human_label") or "").strip()
    if label:
        return " ".join(label.split()).upper()
    return str(turn.get("speaker", "SPEAKER_UNKNOWN"))


def split_readable_paragraphs(chunks: list[str]) -> list[str]:
    paragraphs: list[str] = []
    current = ""

    for chunk_index, chunk in enumerate(chunks):
        chunk = " ".join(str(chunk).split())
        if not chunk:
            continue

        if chunk_index > 0 and len(current) >= READABLE_MIN_PARAGRAPH_CHARS:
            paragraphs.append(current)
            current = ""

        for sentence in split_sentences(chunk):
            candidate = append_text(current, sentence)
            if current and len(candidate) > READABLE_HARD_PARAGRAPH_CHARS:
                paragraphs.append(current)
                current = sentence
                continue

            current = candidate
            if len(current) >= READABLE_SOFT_PARAGRAPH_CHARS and sentence_ends(current):
                paragraphs.append(current)
                current = ""

    if current:
        paragraphs.append(current)

    return merge_short_trailing_paragraph(paragraphs)


def split_sentences(text: str) -> list[str]:
    sentences = [part.strip() for part in SENTENCE_BOUNDARY_RE.split(text) if part.strip()]
    return sentences or [text]


def append_text(left: str, right: str) -> str:
    return f"{left} {right}".strip() if left else right.strip()


def sentence_ends(text: str) -> bool:
    return text.rstrip().endswith((".", "!", "?"))


def merge_short_trailing_paragraph(paragraphs: list[str]) -> list[str]:
    if len(paragraphs) >= 2 and len(paragraphs[-1]) < READABLE_MIN_PARAGRAPH_CHARS:
        paragraphs[-2] = append_text(paragraphs[-2], paragraphs[-1])
        paragraphs.pop()
    return paragraphs


def read_speaker_map(path: Path) -> dict[str, dict[str, str]]:
    if not path.exists():
        return {}
    mapping: dict[str, dict[str, str]] = {}
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        for row in csv.DictReader(handle):
            interview_id = row.get("interview_id", "")
            speaker_label = row.get("speaker_label", "")
            human_label = row.get("human_label", "")
            if interview_id and speaker_label and human_label:
                mapping.setdefault(interview_id, {})[speaker_label] = human_label
    return mapping


def write_empty_speaker_map(path: Path) -> None:
    if path.exists():
        return
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=["interview_id", "speaker_label", "human_label", "confidence", "notes"])
        writer.writeheader()


def find_whisperx_json(paths: Paths, interview_id: str) -> Path | None:
    candidates = [
        paths.asr_dir / f"{interview_id}.json",
        paths.asr_dir / "json" / f"{interview_id}.json",
        paths.asr_dir / "json" / f"{interview_id}.whisperx.json",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    for candidate in paths.asr_dir.rglob(f"*{interview_id}*.json"):
        return candidate
    return None


def write_markdown(path: Path, canonical: dict[str, Any]) -> None:
    lines = [f"# {canonical['interview_id']}", ""]
    for block in build_readable_blocks(canonical["turns"]):
        paragraphs = block["paragraphs"]
        if not paragraphs:
            continue

        lines.append(f"[{format_timestamp(block['start'])}] **{block['label']}:** {paragraphs[0]}")
        lines.append("")
        for paragraph in paragraphs[1:]:
            lines.append(paragraph)
            lines.append("")
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def write_nvivo_tsv(path: Path, canonical: dict[str, Any]) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle, delimiter="\t")
        writer.writerow(["Timespan", "Speaker", "Content"])
        for turn in canonical["turns"]:
            span = f"{format_timestamp(turn['start'], millis=True)}-{format_timestamp(turn['end'], millis=True)}"
            flags = " ".join(f"[{FLAG_DISPLAY_LABELS.get(flag, flag)}]" for flag in turn_flags(turn))
            content = append_text(" ".join(str(turn.get("text", "")).split()), flags)
            writer.writerow([span, turn.get("human_label") or turn.get("speaker"), content])


def write_docx_if_available(path: Path, canonical: dict[str, Any]) -> None:
    try:
        from docx import Document
    except ImportError:
        return

    document = Document()
    document.add_heading(str(canonical["interview_id"]), level=1)
    for block in build_readable_blocks(canonical["turns"]):
        paragraphs = block["paragraphs"]
        if not paragraphs:
            continue

        paragraph = document.add_paragraph()
        paragraph.add_run(f"[{format_timestamp(block['start'])}] ")
        paragraph.add_run(f"{block['label']}: ").bold = True
        paragraph.add_run(paragraphs[0])

        for text in paragraphs[1:]:
            document.add_paragraph(text)
    document.save(path)


def write_srt(path: Path, canonical: dict[str, Any]) -> None:
    cues: list[str] = []
    cue_number = 1
    for turn in canonical.get("turns", []):
        text = subtitle_text(turn)
        if not text:
            continue
        start, end = subtitle_bounds(turn)
        cues.extend(
            [
                str(cue_number),
                f"{format_subtitle_timestamp(start, ',')} --> {format_subtitle_timestamp(end, ',')}",
                text,
                "",
            ]
        )
        cue_number += 1
    path.write_text("\n".join(cues).rstrip() + "\n", encoding="utf-8")


def write_vtt(path: Path, canonical: dict[str, Any]) -> None:
    cues = ["WEBVTT", ""]
    for turn in canonical.get("turns", []):
        text = subtitle_text(turn)
        if not text:
            continue
        start, end = subtitle_bounds(turn)
        cues.extend(
            [
                f"{format_subtitle_timestamp(start, '.')} --> {format_subtitle_timestamp(end, '.')}",
                text,
                "",
            ]
        )
    path.write_text("\n".join(cues).rstrip() + "\n", encoding="utf-8")


def write_turns_csv(path: Path, canonical: dict[str, Any]) -> None:
    write_turns_delimited(path, canonical, delimiter=",")


def write_turns_tsv(path: Path, canonical: dict[str, Any]) -> None:
    write_turns_delimited(path, canonical, delimiter="\t")


def write_turns_delimited(path: Path, canonical: dict[str, Any], delimiter: str) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "interview_id",
                "turn_id",
                "start",
                "end",
                "start_timestamp",
                "end_timestamp",
                "speaker",
                "human_label",
                "flags",
                "notes",
                "text",
            ],
            delimiter=delimiter,
        )
        writer.writeheader()
        for index, turn in enumerate(canonical.get("turns", [])):
            start, end = subtitle_bounds(turn)
            writer.writerow(
                {
                    "interview_id": canonical.get("interview_id", ""),
                    "turn_id": turn.get("id") or f"turn_{index + 1:06d}",
                    "start": f"{start:.3f}",
                    "end": f"{end:.3f}",
                    "start_timestamp": format_timestamp(start, millis=True),
                    "end_timestamp": format_timestamp(end, millis=True),
                    "speaker": turn.get("speaker", ""),
                    "human_label": turn.get("human_label", ""),
                    "flags": ";".join(turn_flags(turn)),
                    "notes": turn.get("notes", ""),
                    "text": " ".join(str(turn.get("text", "")).split()),
                }
            )


def subtitle_text(turn: dict[str, Any]) -> str:
    text = " ".join(str(turn.get("text", "")).split())
    if not text and not turn_flags(turn):
        return ""
    label = format_speaker_label(str(turn.get("human_label") or turn.get("speaker") or "Falante"))
    flags = " ".join(f"[{FLAG_DISPLAY_LABELS.get(flag, flag)}]" for flag in turn_flags(turn))
    body = append_text(text, flags)
    return f"{label}: {body}".strip()


def turn_flags(turn: dict[str, Any]) -> list[str]:
    flags = turn.get("flags", [])
    if not isinstance(flags, list):
        return []
    return [str(flag).strip() for flag in flags if str(flag).strip()]


def subtitle_bounds(turn: dict[str, Any]) -> tuple[float, float]:
    start = float(turn.get("start", 0) or 0)
    end = float(turn.get("end", start) or start)
    if end <= start:
        end = start + 0.25
    return start, end


def format_subtitle_timestamp(seconds: float, millisecond_separator: str) -> str:
    return format_timestamp(seconds, millis=True).replace(".", millisecond_separator)
